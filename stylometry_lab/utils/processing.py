"""
utils/processing.py — Text Processing Utilities

This module handles text loading, cleaning, tokenization, and chunking.
All functions are type-hinted and include comprehensive docstrings.
"""

import re
import logging
from typing import List, Optional, Union, BinaryIO

logger = logging.getLogger(__name__)

# ============================================================================
# IMPORTS & SETUP
# ============================================================================

try:
    import nltk
    from nltk import word_tokenize, sent_tokenize
    nltk.download("punkt", quiet=True)
    HAVE_NLTK = True
except Exception:
    HAVE_NLTK = False

try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    HAVE_SPACY = True
except Exception:
    HAVE_SPACY = False

# Regex patterns
_NONALPHA_RE = re.compile(r"[^a-zA-Zà-žÀ-Ž0-9'\s]")
_WHITESPACE_RE = re.compile(r"\s+")


# ============================================================================
# FILE LOADING - SUPPORTS TXT, PDF, DOCX
# ============================================================================

def load_txt_file(
    uploaded_file: Optional[BinaryIO],
    filename: Optional[str] = None
) -> str:
    """
    Load text from uploaded file, handling multiple file types and encodings.

    Automatically detects file type (txt, pdf, docx) from filename extension
    and uses appropriate parsing method.

    Args:
        uploaded_file: File object from streamlit uploader
        filename: Optional filename to detect file type from extension

    Returns:
        Text content or empty string if file is None/empty

    Example:
        >>> uploaded = st.file_uploader("Upload text")
        >>> text = load_txt_file(uploaded, uploaded.name)
    """
    if uploaded_file is None:
        return ""

    try:
        raw = uploaded_file.read()
        if raw is None or len(raw) == 0:
            return ""

        # Detect file type from filename if provided
        if filename:
            filename_lower = filename.lower()

            if filename_lower.endswith('.pdf'):
                return _extract_text_from_pdf(raw)
            elif filename_lower.endswith('.docx'):
                return _extract_text_from_docx(raw)

        # Default: treat as text file
        return _extract_text_from_bytes(raw)

    except Exception as e:
        logger.error(f"Failed to load text file: {e}")
        return ""


def _extract_text_from_bytes(data: bytes) -> str:
    """Extract text from raw bytes (text file format)."""
    if isinstance(data, str):
        return data

    # Try different encodings
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

    for encoding in encodings:
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, AttributeError):
            continue

    # Last resort
    return data.decode('utf-8', errors='ignore')


def _extract_text_from_pdf(data: bytes) -> str:
    """Extract text from PDF file."""
    try:
        import pdfplumber
        from io import BytesIO

        pdf_file = BytesIO(data)
        text = ""

        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        return text.strip()

    except ImportError:
        logger.error("pdfplumber not installed. Run: pip install pdfplumber")
        return ""
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""


def _extract_text_from_docx(data: bytes) -> str:
    """Extract text from DOCX (Microsoft Word) file."""
    try:
        from docx import Document
        from io import BytesIO

        docx_file = BytesIO(data)
        doc = Document(docx_file)

        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        return text.strip()

    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


# Alias for backward compatibility
extract_text_file = load_txt_file


# ============================================================================
# TEXT CLEANING
# ============================================================================

def clean_text(text: str) -> str:
    """
    Remove special characters and normalize whitespace.

    Performs:
    1. Normalizes line endings
    2. Converts to lowercase
    3. Removes special characters (keeps letters, numbers, spaces)
    4. Collapses multiple spaces

    Args:
        text: Raw text to clean

    Returns:
        Cleaned text

    Example:
        >>> clean_text("Hello!!!  WORLD???")
        'hello world'
    """
    if not text:
        return ""

    # Normalize line endings
    t = text.replace("\r\n", " ").replace("\n", " ")

    # Lowercase
    t = t.lower()

    # Remove special characters
    t = re.sub(_NONALPHA_RE, " ", t)

    # Normalize whitespace
    t = re.sub(r"\s+", " ", t).strip()

    return t


def clean_text_basic(text: str) -> str:
    """
    Basic text cleaning without lowercasing.

    Only normalizes whitespace and line endings.
    Preserves original case and most punctuation.

    Args:
        text: Raw text to clean

    Returns:
        Cleaned text with preserved case

    Example:
        >>> clean_text_basic("Hello\\n\\n  WORLD")
        'Hello WORLD'
    """
    if not text:
        return ""

    # Normalize line endings
    t = text.replace("\r\n", " ").replace("\n", " ")

    # Normalize whitespace
    t = re.sub(r"\s+", " ", t).strip()

    return t


# ============================================================================
# TOKENIZATION
# ============================================================================

def clean_and_tokenize(text: str) -> List[str]:
    """
    Clean and tokenize text using NLTK or simple split.

    Args:
        text: Raw text to tokenize

    Returns:
        List of lowercased tokens

    Example:
        >>> clean_and_tokenize("Hello, World!")
        ['hello', 'world']
    """
    if not text:
        return []

    t = clean_text(text)

    try:
        if HAVE_NLTK:
            tokens = word_tokenize(t)
        else:
            tokens = t.split()

        # Filter short tokens
        tokens = [tok for tok in tokens if len(tok) >= 2]
        return tokens

    except Exception as e:
        logger.error(f"Tokenization failed: {e}")
        return t.split()


def tokenize_with_choice(
    text: str,
    choice: str
) -> List[str]:
    """
    Tokenize text using multiple available tokenization strategies.

    Supports multiple tokenization methods and tries best available:
    1. "whitespace" - Simple whitespace split
    2. "regex" - Regex-based tokenization (handles contractions)
    3. "spacy" - Advanced spacy tokenization (best quality)

    Args:
        text: Input text to tokenize
        choice: Tokenizer choice as string. Options:
                - "whitespace": Simple split() method (fastest)
                - "regex": Regex pattern matching (good for punctuation)
                - "spacy": Spacy tokenizer (slowest but best quality)
                - Can combine: "spacy nltk" (tries first, falls back to second)

    Returns:
        List of tokens (strings) in lowercase
        Empty list if text is None or empty

    Example:
        >>> text = "Don't worry, it's fine!"
        >>> tokenize_with_choice(text, "whitespace")
        ['don't', 'worry,', 'it's', 'fine!']

        >>> tokenize_with_choice(text, "regex")
        ['don't', 'worry', 'it's', 'fine']

        >>> tokenize_with_choice(text, "spacy")
        ['do', 'n't', 'worry', ',', 'it', ''s', 'fine', '!']

    Notes:
        - Returns empty list if text is None
        - All tokens converted to lowercase
        - Spacy tries models: sm → md → lg → blank
        - Falls back to next method if preferred unavailable
    """
    if text is None or not text.strip():
        return []

    choice = (choice or "").lower()

    # Try whitespace first (fastest)
    if "whitespace" in choice:
        tokens = [t for t in text.split() if t.strip()]
        logger.debug(f"Tokenized with whitespace: {len(tokens)} tokens")
        return tokens

    # Try regex (balanced)
    if "regex" in choice:
        tokens = re.findall(r"\w+['-]?\w*|\w+", text.lower())
        logger.debug(f"Tokenized with regex: {len(tokens)} tokens")
        return tokens

    # Try spacy (best quality)
    if "spacy" in choice:
        try:
            if not HAVE_SPACY:
                logger.warning("Spacy not available, falling back to regex")
                return tokenize_with_choice(text, "regex")

            doc = nlp(text)
            tokens = [tok.text.lower() for tok in doc]
            logger.debug(f"Tokenized with spacy: {len(tokens)} tokens")
            return tokens

        except Exception as e:
            logger.warning(f"Spacy tokenization failed: {e}, falling back to regex")
            return tokenize_with_choice(text, "regex")

    # Default fallback to whitespace
    logger.debug(f"Unknown choice '{choice}', using whitespace")
    return tokenize_with_choice(text, "whitespace")


# ============================================================================
# CHUNKING
# ============================================================================

def chunk_words(
    tokens: List[str],
    size: int = 2000
) -> List[List[str]]:
    """
    Split token list into fixed-size chunks.

    Args:
        tokens: List of tokens to chunk
        size: Chunk size (default 2000 tokens)

    Returns:
        List of token chunks

    Example:
        >>> tokens = ['the', 'quick', 'brown', 'fox', ...]
        >>> chunks = chunk_words(tokens, size=3)
        >>> len(chunks)
        2
        >>> chunks[0]
        ['the', 'quick', 'brown']
    """
    if not tokens:
        return []

    chunks = []
    for i in range(0, len(tokens), size):
        chunk = tokens[i:i + size]
        if chunk:
            chunks.append(chunk)

    logger.debug(f"Created {len(chunks)} chunks of size ~{size}")
    return chunks


def chunk_sentences(
    text: str,
    sentences_per_chunk: int = 5
) -> List[List[str]]:
    """
    Split text into chunks based on sentence count.

    Args:
        text: Input text to chunk
        sentences_per_chunk: Number of sentences per chunk

    Returns:
        List of token chunks

    Example:
        >>> text = "Sentence one. Sentence two. Sentence three. Four. Five. Six."
        >>> chunks = chunk_sentences(text, sentences_per_chunk=2)
        >>> len(chunks)
        3
    """
    if not HAVE_NLTK:
        logger.warning("NLTK not available, falling back to word chunking")
        tokens = clean_and_tokenize(text)
        return chunk_words(tokens, 100)

    try:
        sentences = sent_tokenize(text)
        chunks = []

        for i in range(0, len(sentences), sentences_per_chunk):
            sentence_group = sentences[i:i + sentences_per_chunk]
            text_chunk = " ".join(sentence_group)
            tokens = clean_and_tokenize(text_chunk)
            if tokens:
                chunks.append(tokens)

        logger.debug(f"Created {len(chunks)} chunks from {len(sentences)} sentences")
        return chunks

    except Exception as e:
        logger.error(f"Sentence chunking failed: {e}, falling back to word chunking")
        tokens = clean_and_tokenize(text)
        return chunk_words(tokens, 100)
import re
from typing import List
try:
    import nltk
    from nltk import word_tokenize, sent_tokenize
    nltk.download("punkt", quiet=True)
    HAVE_NLTK = True
except Exception:
    HAVE_NLTK = False

def load_txt_file(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    raw = uploaded_file.read()
    if raw is None:
        return ""
    if isinstance(raw, (bytes, bytearray)):
        try:
            return raw.decode("utf-8")
        except Exception:
            try:
                return raw.decode("latin-1")
            except Exception:
                return raw.decode("utf-8", errors="ignore")
    return str(raw)

_nonalpha_re = re.compile(r"[^a-zA-Zà-žÀ-Ž0-9'\s]")

def clean_text(text: str) -> str:
    if not text:
        return ""
    t = text.replace("\r\n", " ").replace("\n", " ")
    t = t.lower()
    t = re.sub(_nonalpha_re, " ", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t

def clean_text_basic(text: str) -> str:
    if not text:
        return ""
    t = text.replace("\r\n", " ").replace("\n", " ")
    t = re.sub(r"\s+", " ", t).strip()
    return t

def clean_and_tokenize(text: str) -> List[str]:
    if not text:
        return []
    t = clean_text(text)
    try:
        if HAVE_NLTK:
            toks = word_tokenize(t)
        else:
            toks = re.findall(r"\w+['-]?\w*|\w+", t)
    except Exception:
        toks = re.findall(r"\w+['-]?\w*|\w+", t)
    toks = [tok for tok in toks if re.search(r"\w", tok)]
    return toks

def tokenize_with_choice(text: str, choice: str):
    if text is None:
        return []
    choice = (choice or "").lower()
    if "whitespace" in choice:
        return [t for t in text.split() if t.strip()]
    if "regex" in choice:
        return re.findall(r"\w+['-]?\w*|\w+", text.lower())
    if "spacy" in choice:
        try:
            import spacy
            for model in ("en_core_web_sm", "en_core_web_md", "en_core_web_lg"):
                try:
                    nlp = spacy.load(model)
                    doc = nlp(text)
                    return [tok.text for tok in doc]
                except Exception:
                    continue
            nlp = spacy.blank("en")
            doc = nlp(text)
            return [tok.text for tok in doc]
        except Exception:
            pass
    if "unicode" in choice:
            return re.findall(r"\p{L}+", text.lower())
    if "char_ngrams" in choice:
        text = re.sub(r"\s+", " ", text.lower())
        n = 3
        return [text[i:i+n] for i in range(len(text)-n+1)]

    
    return clean_and_tokenize(text)

def chunk_words(tokens: List[str], size: int = 2000) -> List[List[str]]:
    if not tokens:
        return []
    return [tokens[i:i+size] for i in range(0, len(tokens), size)]

def chunk_text(text: str, chunk_size: int = 2000) -> List[str]:
    if not text:
        return []
    try:
        toks = word_tokenize(text) if HAVE_NLTK else re.findall(r"\w+['-]?\w*|\w+", text.lower())
    except Exception:
        toks = re.findall(r"\w+['-]?\w*|\w+", text.lower())
    chunks = []
    for i in range(0, len(toks), chunk_size):
        chunk = " ".join(toks[i:i+chunk_size])
        chunks.append(chunk)
    return chunks
# -------------------------------
# DOCUMENT EXTRACTORS
# -------------------------------
import pdfplumber
from docx import Document

def extract_text_from_pdf(file_obj):
    text = []
    with pdfplumber.open(file_obj) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text)

def extract_text_from_docx(file_obj):
    doc = Document(file_obj)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs)

def extract_text_file(file_obj, filename):
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_obj)
    elif name.endswith(".docx"):
        return extract_text_from_docx(file_obj)
    elif name.endswith(".txt"):
        return file_obj.read().decode("utf-8")
    else:
        raise ValueError("Unsupported file format: only PDF, DOCX, TXT allowed")
